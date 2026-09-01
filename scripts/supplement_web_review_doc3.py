from docx import Document
from docx.shared import Pt

SOURCE = r"C:\Users\weimo\Downloads\具身智能社团官网——代码架构与设计 (3).docx"
OUTPUT = r"C:\Users\weimo\Downloads\具身智能社团官网——代码架构与设计 (3)-网页技术要点补充版.docx"

doc = Document(SOURCE)
doc.add_page_break()

def add_marked(text, level=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.underline = True
    if level:
        run.bold = True
        run.font.size = Pt(14 if level == 1 else 12)
    return p

add_marked("五、现有网页值得补充的技术要点", level=1)
add_marked("以下内容是结合当前官网页面、交互代码和接口实现检查后补充的网页工程要点。所有新增内容使用下划线标记。")

sections = [
    ("5.1 前端请求超时与取消", "新闻、荣誉、联系表单和字幕翻译都通过 fetch 请求数据。建议为前端请求增加 AbortController 和明确超时，例如表单请求超过 10 秒自动结束并恢复提交按钮；页面切换时取消不再需要的请求，避免旧页面请求返回后继续更新已离开的 DOM。服务端调用飞书、微信公众号和 AI 翻译服务时也应设置连接超时和响应超时。"),
    ("5.2 React 错误边界与接口错误兜底", "当 React 渲染、页面模板或专项模块发生异常时，应显示可理解的错误兜底界面，而不是让页面区域空白。接口失败时应区分网络错误、缓存为空、服务暂不可用和表单校验失败，并提供重试、返回首页或保留旧内容等操作。"),
    ("5.3 未知路由和 404 页面", "当前路由匹配不到页面时不应静默回到首页。建议为未知 Hash 路由显示 404 页面，提供返回首页和联系我们的入口，并在开发日志中记录非法路由，便于发现错误链接。"),
    ("5.4 动态 HTML 与 XSS 防护", "项目使用 dangerouslySetInnerHTML 生成页面模板。当前模板主要来自代码，但新闻、荣誉、标题、链接和图片地址来自缓存或外部平台，后续改动时必须避免把未经处理的数据直接拼接进 HTML。普通文本优先使用 textContent；必须渲染 HTML 时使用白名单清洗；外部链接只允许 HTTPS，图片 URL 进行协议、域名和路径校验。"),
    ("5.5 浏览器安全响应头", "除了已有的 X-Content-Type-Options，还建议配置 Content-Security-Policy、Strict-Transport-Security、Referrer-Policy、Permissions-Policy 和 frame-ancestors。Leaflet、地图瓦片和翻译服务需要在 CSP 中明确允许的域名，避免第三方脚本被任意替换或扩展。"),
    ("5.6 移动端导航与键盘焦点", "移动端菜单应支持键盘和屏幕阅读器操作：菜单打开后焦点进入菜单，按 Esc 可以关闭，关闭后焦点回到菜单按钮，菜单打开期间背景内容不应继续被键盘访问。弹窗和视频预览也应遵循相同的焦点进入、焦点锁定和焦点恢复规则。"),
    ("5.7 轮播、Tab 和视频的无障碍补充", "现有页面已经使用 aria-label、aria-selected 和 aria-live，建议继续补充 Tab 与面板之间的 aria-controls，隐藏非活动面板，给轮播提供明确的暂停按钮，并保证字幕、播放状态和加载失败状态不仅通过颜色或动画表达。"),
    ("5.8 路由级 SEO", "当前入口页面使用统一的 title 和 description。建议在路由或语言切换时同步更新页面标题、描述、canonical、Open Graph 标题和图片；产品页、新闻页、联系页应使用不同的 SEO 文案，并补充 robots.txt、sitemap.xml 以及公司、产品和新闻的结构化数据。"),
    ("5.9 图片、视频和第三方资源降级", "新闻、荣誉、地图、视频和 CDN 资源都需要失败兜底。图片应有占位图和加载失败处理；地图加载失败时显示地址和外部地图链接；Leaflet CDN 失败时页面仍能使用；视频无法播放时显示文字说明或下载入口。第三方 CDN 应固定版本，并评估 SRI 或本地备用资源。"),
    ("5.10 轮询与页面可见性", "荣誉模块当前存在定时刷新。建议页面不可见时暂停轮询，重新可见时再刷新；同步失败后采用递增退避；内容没有变化时不要重复替换 DOM。这样可以减少后台标签页的网络请求和重复渲染。"),
    ("5.11 缓存验证与性能预算", "新闻和荣誉接口可以考虑 ETag、Last-Modified 或短时缓存，内容未变化时返回 304。还应设定首屏 JavaScript、图片大小、首屏加载时间和 Core Web Vitals 目标；非首屏图片使用懒加载，视频不自动下载全部内容，首屏资源优先使用压缩格式。"),
    ("5.12 监控、请求追踪与隐私脱敏", "建议为接口请求和后台同步任务增加 requestId、耗时、状态码、缓存更新时间和外部服务错误指标，并提供健康检查和就绪检查接口。日志不得记录 Token、手机号、邮箱和完整留言内容；联系表单应明确隐私政策、保存期限和访问权限。"),
    ("5.13 真实设备测试与发布检查", "除现有交互冒烟、视频 Range 和视觉回归测试外，还应覆盖 iPhone Safari、Android Chrome、窄屏设备、断网、图片失败、地图失败、翻译超时、表单重复提交和所有页面路由。发布前至少执行构建检查、接口契约检查、可访问性检查和关键页面截图回归。"),
    ("5.14 优雅关闭与静态资源版本", "Node 服务收到退出信号时，应停止接收新请求、等待当前请求和同步任务结束、清理定时器后再退出。生产构建的 JS、CSS 和图片建议使用内容哈希文件名，配合长期缓存，避免发布后用户继续使用旧资源。"),
]

for heading, body in sections:
    add_marked(heading, level=2)
    add_marked(body)

add_marked("网页检查结论：当前项目的页面结构、路由切换、交互清理、接口缓存和基础防刷已经具备较好的基础。下一阶段优先补充请求超时与取消、404 和错误兜底、XSS 防护、SEO、移动端无障碍、第三方资源降级以及真实设备测试。")

doc.save(OUTPUT)
print(OUTPUT)
