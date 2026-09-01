from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SOURCE = r"C:\Users\weimo\Downloads\具身智能社团官网——代码架构与设计 (3)-按原格式补充版.docx"
OUTPUT = r"C:\Users\weimo\Downloads\具身智能社团官网——代码架构与设计 (3)-按原格式代码块版.docx"

doc = Document(SOURCE)

# Remove the underline used as the previous change marker.
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.underline = False
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.underline = False

code_blocks = [
    """const NEWS_CACHE_TTL = 60_000;
let newsCache = null;
let newsRequest = null;

function loadNewsData({ force = false } = {}) {
  const fresh = newsCache &&
    Date.now() - newsCache.loadedAt < NEWS_CACHE_TTL;
  if (!force && fresh) return Promise.resolve(newsCache.payload);
  if (newsRequest) return newsRequest;
  newsRequest = fetch(\"/api/news\")
    .then((response) => response.ok ? response.json() : null)
    .finally(() => { newsRequest = null; });
  return newsRequest;
}""",
    """const storageKey = `oculotronics:scroll:${page}`;

const savePosition = () => {
  try { sessionStorage.setItem(storageKey, String(window.scrollY)); } catch {}
};

window.addEventListener(\"pagehide\", savePosition);
window.addEventListener(\"oculotronics:content-ready\", onContentReady);

requestAnimationFrame(() =>
  requestAnimationFrame(() => window.scrollTo(0, position))
);""",
    """const reducedMotion =
  matchMedia(\"(prefers-reduced-motion: reduce)\").matches;

const observer = new IntersectionObserver((entries) => {
  entries.forEach((entry) => {
    if (entry.isIntersecting) entry.target.classList.add(\"is-visible\");
  });
});

return () => {
  observer.disconnect();
  splitTextObserver?.disconnect();
};""",
    """const currentTime = video.currentTime;
const resume = !video.paused;
video.src = source.src;
video.addEventListener(\"loadedmetadata\", () => {
  video.currentTime = Math.min(currentTime, video.duration || currentTime);
  if (resume) video.play().catch(() => {});
}, { once: true });

const key = `${targetLanguage}:${sourceText}`;
const cached = translationCache.get(key);
if (cached) return cached;
// translationCache 的条目数限制为 300，避免无限增长。""",
]

targets = [p for p in doc.paragraphs if p.text.strip().startswith("代码解释：")]
if len(targets) < 4:
    raise RuntimeError(f"Expected four new code explanation paragraphs, found {len(targets)}")

def shade_cell(cell, fill="F3F5F7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def add_code_table_after(paragraph, code):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    shade_cell(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(code)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(40, 48, 55)
    # Move the newly-created table directly after the explanation paragraph.
    paragraph._p.addnext(table._tbl)

for paragraph, code in zip(targets[-4:], code_blocks):
    add_code_table_after(paragraph, code)

doc.save(OUTPUT)
print(OUTPUT)
