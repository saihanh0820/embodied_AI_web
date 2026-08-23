from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = 'weimou_web——Coding Review.docx'
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(.7); sec.bottom_margin = Inches(.7); sec.left_margin = Inches(.8); sec.right_margin = Inches(.8)
styles = doc.styles
styles['Normal'].font.name = '微软雅黑'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑'); styles['Normal'].font.size = Pt(10.5)
for name, size, color in [('Title', 24, '163A5F'), ('Heading 1', 15, '163A5F'), ('Heading 2', 12, '2F6B8A')]:
    s=styles[name]; s.font.name='微软雅黑'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def table(headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c,'DCEAF3'); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: r.font.bold=True
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): cells[i].text=str(v); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph()
    return t

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('weimou_web 项目 Coding Review'); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=RGBColor(22,58,95)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('代码架构、可靠性与测试质量评审稿').italic=True
doc.add_paragraph('评审日期：2026-08-19    评审范围：React/Vite 前端、Node HTTP 服务、交互冒烟测试与构建链路')
doc.add_heading('一、评审结论', level=1)
doc.add_paragraph('当前项目未发现阻断发布的 P0/P1 级问题。代码分层清晰，缓存优先策略、服务端校验与页面生命周期清理方向正确。评审期间已修复 4 项会影响安全性、长期稳定性或测试可信度的问题；构建与 Range 接口验证通过。交互冒烟脚本仍有一处与当前新闻卡片结构不一致的历史断言，已定位为测试维护项，建议合并前完成脚本重整。')
doc.add_heading('二、评审依据', level=1)
table(['验证项','结果','证据/说明'],[
    ('生产构建','通过','npm run build；Vite 构建成功，仅有 modernTargets 配置提示'),
    ('媒体 Range 接口','通过','npm run test:range；返回 206、Content-Range 正确、长度 1024'),
    ('交互冒烟','待修整','已补齐 API 服务启动，但脚本仍依赖已移除的新闻图片/年份筛选结构'),
    ('架构文档一致性','基本一致','源文档中的分层、缓存、校验与生命周期描述与代码主路径相符')
], [1.4,1.0,4.6])
doc.add_heading('三、已发现并完成的改进', level=1)
table(['编号','优先级','问题','处理结果'],[
    ('CR-01','高','翻译限流直接信任原始 X-Forwarded-For，可被伪造来源地址绕过限流','改为统一使用 getClientAddress，并遵循 TRUST_PROXY 配置；未启用可信代理时不采信伪造头'),
    ('CR-02','中','translationRequests 仅增不减，长时间运行会造成进程内存持续增长','新增每分钟清理任务，删除过期客户端记录，并使用 unref 避免阻止进程退出'),
    ('CR-03','中','交互冒烟测试只启动 Vite，/api/news 会被代理到错误目标，动态新闻场景不可重复','测试启动时分配 API_PORT、拉起 server.js，并在执行前等待 API 可用'),
    ('CR-04','中','测试包含已移除的主题切换按钮断言，且把产品页有意差异误判为结构回归','移除过时主题断言；产品联系带比较改为比较稳定结构、文案、表面与核心几何')
], [0.65,0.65,3.0,3.0])
doc.add_heading('四、建议继续改进', level=1)
for text in [
    '将翻译/线索限流从单进程 Map 提升到网关、WAF 或 Redis；多实例部署时，当前实现无法形成全局限流。',
    '为 Feishu、翻译、微信等出站 fetch 增加 AbortSignal.timeout，并记录超时、重试和上游错误指标，避免慢请求占用连接。',
    '补充服务端接口测试：恶意/缺失字段、代理头处理、限流窗口、Feishu 事件签名与缓存回退。',
    '继续收敛远程内容渲染边界：默认使用 textContent；若未来引入 HTML 内容，应统一经过白名单 sanitizer。',
    '明确 Vite legacy 浏览器支持矩阵，处理 modernTargets 覆盖内置 targets 的构建提示。'
]: doc.add_paragraph(text, style='List Bullet')
doc.add_heading('五、合并前验收清单', level=1)
table(['检查项','责任建议','状态'],[
    ('npm run build','前端/CI','已通过'),('npm run test:range','服务端/CI','已通过'),('npm run test:interactions','前端/QA','修整新闻数据断言后重跑'),('接口异常与限流测试','后端/QA','建议新增'),('生产代理与多实例限流方案','架构/运维','上线前确认')
], [3.2,2.0,1.4])
doc.add_heading('六、评审意见', level=1)
doc.add_paragraph('建议按 CR-01～CR-04 的实现合入当前分支，并把交互冒烟脚本的新闻场景改为以当前 DOM 契约为准。完成出站超时、接口异常测试和生产限流方案后，项目可进入常规发布评审。')
doc.add_paragraph('本稿为代码评审记录，不替代安全渗透测试、浏览器兼容性矩阵或生产容量压测。')
doc.save(OUT)
print(OUT)
