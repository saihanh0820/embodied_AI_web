from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "weimou_web_软件与底层架构审查.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F1F1F"
MUTED = "666666"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "B7C3D0"


def set_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for index, width in enumerate(widths):
        grid.gridCol_lst[index].set(qn("w:w"), str(width))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_table_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), BORDER)
        borders.append(tag)
    table._tbl.tblPr.append(borders)


def add_run(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_font(run, **kwargs)
    return run


def add_body(doc, text, after=6):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(after)
    add_run(p, text, size=11, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    add_run(p, text, size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_run(p, header, size=10, color=INK, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_run(p, value, size=10.5, color=INK, bold=index == 0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, f"{title}  ", size=10.5, color=DARK_BLUE, bold=True)
    add_run(p, text, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_footer_field(paragraph):
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for level, size, color, before, after in ((1, 16, BLUE, 16, 8), (2, 13, BLUE, 12, 6), (3, 12, DARK_BLUE, 8, 4)):
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    add_run(header, "weimou_web | 软件与底层架构审查", size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    add_run(footer, "内部审查资料 | 2026-08-14 | 第 ", size=9, color=MUTED)
    add_footer_field(footer)
    add_run(footer, " 页", size=9, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(4)
    add_run(title, "软件与底层架构审查", size=23, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    add_run(subtitle, "weimou_web 项目技术资产盘点", size=14, color=MUTED)

    metadata = [("审查对象", "weimou_web 项目"), ("审查范围", "源码、构建配置、环境变量示例与本地脚本"), ("审查日期", "2026-08-14"), ("结论", "轻量全栈官网：React 前端 + Node.js 原生 HTTP 服务")]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_run(p, f"{label}：", size=11, color=INK, bold=True)
        add_run(p, value, size=11, color=INK)

    add_heading(doc, "一、执行摘要", 1)
    add_body(doc, "本项目是以 React 为浏览器端运行时、Vite 为构建工具、Node.js 原生 HTTP 服务为 API 层的轻量全栈官网。官网内容和线索数据通过飞书开放平台与飞书多维表格管理；新闻、荣誉及其图片会同步到本地文件缓存，保证外部平台暂时不可用时的网站读取能力。")
    add_callout(doc, "审查结论", "项目未引入数据库、ORM、Express/Koa、TypeScript、Redux、UI 组件库、Docker、云函数、消息队列、监控平台或支付 SDK。")

    add_heading(doc, "二、核心软件与开发工具", 1)
    add_table(doc, ["类别", "软件 / 版本", "用途"], [
        ("前端运行时", "React 19.2.4；React DOM 19.2.4", "页面挂载、路由变化后的组件生命周期管理。"),
        ("构建与压缩", "Vite 7.3.6；esbuild；Terser", "本地开发服务器、生产构建与资源压缩。"),
        ("兼容层", "@vitejs/plugin-legacy；Babel；Browserslist", "生成旧浏览器兼容代码及必要 polyfill。"),
        ("后端运行时", "Node.js 核心模块（http、fs、path、crypto）", "API、静态文件托管、文件缓存、哈希去重与请求处理；未使用 Express。"),
        ("验证工具", "Google Chrome Headless；Chrome DevTools Protocol", "交互冒烟、国际化审计和视觉回归测试。"),
        ("开发辅助", "PowerShell；Python", "素材整理、图片处理和联系表等辅助脚本。")
    ], [1700, 3500, 4160])

    add_heading(doc, "三、外部平台与服务", 1)
    add_table(doc, ["服务", "接入范围", "启用状态"], [
        ("飞书开放平台", "租户访问令牌、线索表单、荣誉及新闻同步、文件上传下载。", "已接入"),
        ("飞书多维表格", "销售线索、企业荣誉和新闻内容的后台数据来源。", "已接入"),
        ("腾讯位置服务", "JavaScript API GL、地址原生地理编码、联系页地图与标记。", "已接入"),
        ("微信公众平台 API", "公众号文章及封面同步至飞书新闻表。", "可选，默认关闭"),
        ("AI 翻译 HTTP 服务", "视频字幕翻译；具体服务商通过环境变量指定。", "可选")
    ], [1800, 5000, 2560])

    add_heading(doc, "四、底层架构与数据流", 1)
    add_heading(doc, "1. 浏览器端", 2)
    add_body(doc, "React 应用使用 Hash 路由管理首页、产品、新闻、支持及联系页面。页面主体由项目内的模板函数生成，React 负责挂载、路由切换和副作用清理；中英文切换由项目内置词典处理。腾讯地图 GL SDK 在联系页按需加载。")
    add_heading(doc, "2. 构建与发布", 2)
    add_body(doc, "Vite 将 React、CSS 与静态资源构建为 dist 目录。构建配置会复制受管理的图片与原始素材，并通过 legacy 插件为目标浏览器生成兼容代码。开发环境将 /api 代理到本机 8787 端口。")
    add_heading(doc, "3. 服务端 API", 2)
    add_body(doc, "Node.js 单一 HTTP 服务同时提供静态文件和 API。主要接口包括健康检查、联系表单、字幕翻译、荣誉/新闻读取与图片读取、飞书事件回调，以及可选的微信同步回调与触发接口。")
    add_heading(doc, "4. 内容与缓存", 2)
    add_body(doc, "服务端按配置定时从飞书多维表格同步荣誉和新闻；同步结果保存在 data/honors 与 data/news 下的 JSON 和图片文件中。项目未配置关系型、文档型或云托管数据库。")

    add_heading(doc, "五、安全与运行治理", 1)
    add_table(doc, ["控制项", "当前实现"], [
        ("凭证管理", "飞书、腾讯地图、微信和翻译服务凭证均由环境变量读取；示例配置不包含真实密钥。"),
        ("表单保护", "限制请求体大小、提交频率、突发请求和单客户端并发数；使用蜜罐字段、幂等键和请求内容摘要避免重复提交。"),
        ("边界校验", "静态文件路径、缓存图片文件名和 JSON 请求体均有校验；API 响应包含 nosniff 头。"),
        ("可用性", "新闻与荣誉使用本地缓存；飞书事件回调不可用或延迟时，可通过定时轮询更新。"),
        ("可选集成", "微信同步默认关闭；字幕翻译需显式配置服务 URL 和令牌后才可调用。")
    ], [2500, 6860])

    add_heading(doc, "六、审查边界与后续建议", 1)
    add_body(doc, "本清单基于当前代码仓库的 package.json、vite.config.js、server.js、src/main.jsx、腾讯地图模块、环境变量示例和 scripts 目录整理。它反映项目已声明或已调用的软件与服务，不等同于生产部署环境中由运维平台额外提供的操作系统、Web 服务器、CDN、证书、日志或监控能力。")
    add_body(doc, "建议在上线审查中另行补充：生产部署位置、域名与 HTTPS 证书管理方式、反向代理与 CDN 配置、备份保留周期、日志审计策略，以及飞书/微信/腾讯地图/翻译服务的账号归属与密钥轮换负责人。")

    doc.core_properties.title = "weimou_web 软件与底层架构审查"
    doc.core_properties.subject = "软件资产与架构审查"
    doc.core_properties.author = "微眸医疗"
    doc.core_properties.comments = "Generated from project source review."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
