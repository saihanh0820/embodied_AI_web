import glob, shutil
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

src = [x for x in glob.glob(r'C:\Users\weimo\Downloads\*.docx') if '具身智能社团官网' in x][0]
out = '具身智能社团官网——代码架构与设计（展示修订版）.docx'
shutil.copyfile(src, out)
d = Document(out)
changes = {
    3: '本项目采用“React 应用壳层 + app.js 命令式 DOM 页面层 + 专项能力模块 + Node.js 服务端 + 本地数据缓存”的混合分层架构。src/main.jsx 负责 React 应用加载、Hash 路由、语言切换和页面生命周期；app.js 负责页面模板、交互逻辑和 DOM 更新；i18n.js 与 styles.css 负责国际化和视觉样式；视频、字幕、地图等能力由独立模块封装；server.js 统一提供静态资源和业务 API，承担输入校验、限流、安全边界与缓存同步。',
    16: '页面层不是纯 React 组件渲染，而是 React 壳层与 app.js 命令式 DOM 页面层的混合实现。React 负责应用入口、路由状态和生命周期；app.js 负责生成页面 HTML、绑定交互、请求数据并更新 DOM。这样既保留了 React 对应用状态和生命周期的管理，也降低了静态营销页面整体迁移为纯组件的成本。',
    17: '具体分工是：main.jsx 只负责应用启动、路由和生命周期；app.js 决定渲染什么页面、绑定哪些交互以及如何处理页面数据。页面切换时由 React effect 统一触发资源初始化和清理，页面模块不直接长期持有跨页面资源。'
}
for idx, text in changes.items():
    if idx < len(d.paragraphs):
        p = d.paragraphs[idx]
        p.text = text
        for run in p.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

first = d.paragraphs[0]
first.text = '【展示修订版说明】黄色高亮为本次根据实际代码补正的内容；未高亮部分沿用原架构文档。'
for run in first.runs:
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
d.save(out)
print(out)
