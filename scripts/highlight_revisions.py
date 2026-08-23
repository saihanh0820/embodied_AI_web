from docx import Document
from docx.enum.text import WD_COLOR_INDEX

path = 'weimou_web——架构展示与集成讲解稿.docx'
d = Document(path)
targets = [
    '混合分层架构',
    'app.js 生成页面结构并请求同源 API',
    'App 组件维护当前页面和语言，但它不是所有业务页面的 JSX 渲染器',
    '这里采用命令式 DOM 方式生成和更新新闻页面',
    '同步可以由启动同步、定时轮询或事件触发',
]
for p in d.paragraphs:
    if any(t in p.text for t in targets):
        for run in p.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

for p in d.paragraphs:
    if p.text.startswith('使用方式：'):
        p.text = '使用方式：黄色高亮部分为本次修订内容；先展示架构表，再按本文顺序打开代码文件。'
        for run in p.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        break
d.save(path)
print(path)
