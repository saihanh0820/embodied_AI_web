from pathlib import Path
from docx import Document

for path in Path('.').glob('weimou_web*.docx'):
    d = Document(path)
    for p in d.paragraphs:
        if '具身智能社团官网' in p.text:
            p.text = p.text.replace('具身智能社团官网', 'weimou_web')
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if '具身智能社团官网' in p.text:
                        p.text = p.text.replace('具身智能社团官网', 'weimou_web')
    d.core_properties.title = d.core_properties.title.replace('具身智能社团官网', 'weimou_web') if d.core_properties.title else d.core_properties.title
    d.save(path)
print('updated docx labels')
