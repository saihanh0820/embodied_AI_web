import glob
from docx import Document
p = [x for x in glob.glob(r'C:\Users\weimo\Downloads\*.docx') if '微眸医疗官网' in x][0]
d = Document(p)
print('FILE', p)
for x in d.paragraphs:
    if x.text.strip(): print(x.text)
print('TABLES', len(d.tables))
for t in d.tables:
    print('---TABLE')
    for row in t.rows: print(' | '.join(c.text.replace('\n',' / ') for c in row.cells))
