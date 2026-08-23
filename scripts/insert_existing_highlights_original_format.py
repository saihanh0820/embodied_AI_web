from docx import Document
from docx.shared import Pt

SOURCE = r"C:\Users\weimo\Downloads\微眸医疗官网——代码架构与设计 (3).docx"
OUTPUT = r"C:\Users\weimo\Downloads\微眸医疗官网——代码架构与设计 (3)-按原格式补充版.docx"

doc = Document(SOURCE)
boundary = next((p for p in doc.paragraphs if p.text.strip().startswith("四、维护边界")), None)
if boundary is None:
    raise RuntimeError("Could not locate maintenance boundary paragraph")

content = [
    ("3.5 前端数据请求的缓存与请求合并", "heading"),
    ("新闻数据在浏览器端设置短时 TTL 缓存，并把尚未完成的请求保存为共享 Promise。首页新闻区和新闻中心同时加载时，后续调用复用同一个请求，不会重复访问 /api/news。", "body"),
    ("核心设计：缓存负责减少重复读取，请求合并负责避免并发重复请求；缓存失效后重新请求，失败时不覆盖已有可用数据。", "body"),
    ("代码解释：weimou_web_app.js 中的 newsCache 保存 payload 和 loadedAt，60 秒内直接返回缓存；newsRequest 保存当前请求，存在时直接返回该 Promise。请求结束后在 finally 中清空 newsRequest，保证下一次失效读取可以重新发起。", "body"),
    ("设计要点：", "body"),
    ("（1）多个页面共享同一份新闻请求，减少网络和外部服务压力；", "body"),
    ("（2）缓存时间短，兼顾页面响应速度和内容更新及时性；", "body"),
    ("（3）请求合并与服务端落盘缓存配合，形成浏览器和服务端两级稳定读取。", "body"),
    ("3.6 异步内容加载后的页面位置恢复", "heading"),
    ("新闻等页面内容不是首屏一次性全部生成，而是在接口返回后插入 DOM。项目对每个页面单独保存 sessionStorage 滚动位置，并在异步内容完成、页面高度恢复后再执行滚动恢复。", "body"),
    ("核心设计：路由切换到新页面时回到顶部；浏览器刷新或返回同一页面时恢复该页面的位置；异步内容加载完成后通过自定义事件通知壳层重新校准。", "body"),
    ("代码解释：src/weimou_web_main.jsx 使用 oculotronics:scroll:<page> 作为存储键，在 pagehide 和 effect 清理时保存 window.scrollY。恢复时连续等待两次 requestAnimationFrame；新闻数据加载完成后，weimou_web_app.js 派发 oculotronics:content-ready，壳层收到后再次恢复位置。", "body"),
    ("设计要点：", "body"),
    ("（1）不同页面的滚动位置互不覆盖；", "body"),
    ("（2）避免在异步列表插入前恢复位置造成跳动；", "body"),
    ("（3）页面展示状态和数据加载状态通过事件完成解耦。", "body"),
    ("3.7 动画系统兼顾性能与用户偏好", "heading"),
    ("页面动画不是一次性全部执行，而是通过 IntersectionObserver 在元素进入视口时触发；同时检查 prefers-reduced-motion，尊重用户对减少动态效果的系统设置。", "body"),
    ("核心设计：只观察即将进入视口的元素；动画结束或页面离开时断开观察器；减少动态效果用户跳过文字拆分、转场和指针跟随等非必要动画。", "body"),
    ("代码解释：weimou_web_app.js 创建 IntersectionObserver 观察 reveal、feature-row 和 section-reveal 元素；滚动进度和部分高频视觉变化通过 requestAnimationFrame 合并到下一帧。代码在 cleanup 中调用 observer.disconnect() 和 splitTextObserver.disconnect()。", "body"),
    ("设计要点：", "body"),
    ("（1）动画计算集中在视口附近，降低首屏和滚动时的开销；", "body"),
    ("（2）减少动态效果不会破坏页面信息结构；", "body"),
    ("（3）观察器和动画资源随页面生命周期释放，避免路由切换后继续执行。", "body"),
    ("3.8 视频播放器保持播放连续性并限制翻译缓存", "heading"),
    ("视频播放器支持清晰度切换、字幕翻译、画中画和全屏。切换视频源时会保留当前播放时间和播放状态；字幕翻译结果在浏览器端缓存，并设置条目上限。", "body"),
    ("核心设计：替换 video source 前记录 currentTime 和 paused 状态，等待 loadedmetadata 后恢复；翻译缓存按目标语言和原字幕组成键，超过上限后淘汰旧条目。", "body"),
    ("代码解释：src/weimou_web_pot-player.js 在切换清晰度时记录 currentTime，监听 loadedmetadata 后恢复进度；translationCache 以 targetLanguage:sourceText 为键，并由 MAX_TRANSLATION_CACHE_ENTRIES=300 控制最大数量，防止长时间使用播放器造成内存无限增长。", "body"),
    ("设计要点：", "body"),
    ("（1）清晰度切换不会把用户带回视频开头；", "body"),
    ("（2）相同字幕不会重复请求翻译服务；", "body"),
    ("（3）缓存有明确上限，播放器长期运行时内存占用可控。", "body"),
]

for text, kind in content:
    p = boundary.insert_paragraph_before()
    run = p.add_run(text)
    run.underline = True
    if kind == "heading":
        run.bold = True
        run.font.size = Pt(12)

doc.save(OUTPUT)
print(OUTPUT)
