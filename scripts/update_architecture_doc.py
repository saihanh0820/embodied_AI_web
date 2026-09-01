import glob, shutil
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement

src = [x for x in glob.glob(r'C:\Users\weimo\Downloads\*.docx') if '具身智能社团官网——代码架构与设计 (2)' in x][0]
out = 'weimou_web——代码架构与设计（代码解释高亮版）.docx'
shutil.copyfile(src, out)
d = Document(out)

def highlight(p):
    for r in p.runs:
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW

repls = {
    '具身智能社团官网': 'weimou_web',
    'src/main.jsx': 'src/main.jsx',
    'app.js': 'app.js',
    'i18n.js': 'i18n.js',
    'styles.css': 'styles.css',
    'server.js': 'server.js',
    'src/pot-player.js': 'src/pot-player.js',
    'src/leaflet-map.js': 'src/leaflet-map.js',
}
for p in d.paragraphs:
    old = p.text
    new = old
    for a,b in repls.items(): new = new.replace(a,b)
    if new != old:
        p.text = new; highlight(p)
for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                old = p.text; new = old
                for a,b in repls.items(): new = new.replace(a,b)
                if new != old:
                    p.text = new; highlight(p)

def insert_after_table(table, title, body):
    p = OxmlElement('w:p')
    table._tbl.addnext(p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(p, table._parent)
    para.paragraph_format.space_before = 4
    para.paragraph_format.space_after = 8
    run = para.add_run(f'{title}：{body}')
    run.font.name = '微软雅黑'
    return para

explanations = {
    1: ('代码解释', 'useEffect 在页面或语言发生变化时重新绑定交互、视频播放器和地图。返回函数是清理阶段：先释放地图，再停止播放器，最后移除 DOM 事件。这样可以防止路由切换后事件重复绑定、计时器残留和第三方资源泄漏。'),
    2: ('代码解释', 'handleNews 只读取本地新闻缓存并返回统一 JSON，不在用户请求链路中等待外部平台。cache 为空时返回空数组，页面仍能正常渲染；syncedAt 用于展示或排查数据新鲜度；source 明确数据来自本地缓存。外部同步由服务端后台任务负责。'),
    3: ('代码解释', '表单请求先根据客户端地址执行限流，再读取并校验 JSON。限流失败立即返回 429，避免无效请求继续消耗资源；readJson 限制请求体大小，防止超大 payload；后续还会继续做字段长度、格式、幂等键和重复提交校验。安全边界放在服务端，不能只依赖前端校验。'),
    4: ('代码解释', 'preloadLeafletMapResources 把地图脚本和样式的加载封装成可复用函数，并用 Promise.all 并行加载，减少等待时间。预加载只负责准备资源，真正初始化地图仍在联系页进入时执行；两者分离可以兼顾首屏性能和按需使用。'),
}
for idx, (title, body) in explanations.items():
    insert_after_table(d.tables[idx], title, body)

d.save(out)
print(out)
