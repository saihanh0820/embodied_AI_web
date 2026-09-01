from docx import Document
from docx.shared import Pt

SOURCE = r"C:\Users\weimo\Downloads\具身智能社团官网——代码架构与设计 (3).docx"
OUTPUT = r"C:\Users\weimo\Downloads\具身智能社团官网——代码架构与设计 (3)-现有网页技术亮点补充版.docx"

doc = Document(SOURCE)
doc.add_page_break()

def add_marked(text, level=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.underline = True
    if level:
        run.bold = True
        run.font.size = Pt(14 if level == 1 else 12)
    return p

add_marked("五、现有网页已实现但原文未展开的技术亮点", level=1)
add_marked("以下不是新增建议，而是从当前网页代码中整理出的已实现能力。它们已经在项目中发挥作用，补充到架构文档中便于后续维护人员理解。新增内容统一使用下划线标记。")

sections = [
    ("5.1 按页面保存和恢复滚动位置", "src/main.jsx 使用 sessionStorage 按页面保存滚动位置，例如为每个 page 生成独立的 oc(u)lotronics:scroll:<page> 键。页面初次打开时会恢复位置，路由切换时回到顶部。对于新闻等异步内容，代码通过两次 requestAnimationFrame 等待布局完成，并监听 oculotronics:content-ready 事件后再恢复滚动，避免数据加载前恢复位置导致偏移。"),
    ("5.2 前端新闻请求的 TTL 缓存与请求合并", "app.js 同时维护 newsCache 和 newsRequest：缓存 60 秒内直接复用结果；已有请求尚未完成时，后续调用共享同一个 Promise，而不是重复发起网络请求。这种 single-flight 设计能避免首页新闻区和新闻中心同时加载时产生重复请求。"),
    ("5.3 异步数据返回前的 DOM 生命周期检查", "新闻和首页新闻加载完成后，会检查 newsGrid.isConnected 或 homeNewsList.isConnected，再决定是否更新 DOM。用户在请求返回前切换了页面时，旧页面节点已脱离文档，代码会放弃更新，从而避免异步回调修改旧页面或产生竞态问题。"),
    ("5.4 联系表单幂等键的前端生成", "联系表单提交时优先使用 crypto.randomUUID() 生成请求标识，并通过 X-Idempotency-Key 发送到服务端；不支持 randomUUID 的浏览器还有时间戳和随机数回退方案。这让重复点击、网络重试和用户重新提交可以被服务端区分，和后端的重复提交识别形成完整闭环。"),
    ("5.5 尊重用户的减少动态效果偏好", "页面动画、文字拆分、产品图片跟随指针旋转和新闻转场都会检查 prefers-reduced-motion: reduce。用户在系统中开启减少动态效果后，页面会降低或跳过动画，并将返回顶部等滚动操作改为非平滑滚动，这体现了无障碍和用户偏好的兼容。"),
    ("5.6 IntersectionObserver 驱动的按需入场动画", "页面通过 IntersectionObserver 只在元素进入视口时添加 reveal 状态，而不是页面加载时一次性执行所有动画。文字拆分也使用独立观察器，并在清理阶段 disconnect，减少首屏计算量和长期监听器残留。"),
    ("5.7 轮播的用户行为与页面可见性控制", "首页产品轮播会在鼠标进入、键盘焦点进入时暂停，鼠标离开或焦点离开后恢复；浏览器标签页进入后台时清除定时器，重新可见时再启动。用户手动操作、键盘操作和后台节能都被纳入同一套轮播生命周期。"),
    ("5.8 requestAnimationFrame 合并高频视觉更新", "滚动进度、视频控制条和地图尺寸更新使用 requestAnimationFrame 合并同一帧内的多次变化，避免 scroll、timeupdate 或 resize 事件每触发一次就立即强制布局和重绘，降低高频交互下的卡顿风险。"),
    ("5.9 图片加载性能和内容预加载的分层策略", "普通图片统一使用 loading=lazy 和 decoding=async；荣誉图片在数据到达后通过 Image 对象预加载并使用 Map 去重；地图资源和周边瓦片也在适合的时机预加载。项目将首屏图片、当前视口图片和下一步可能出现的图片区分处理，兼顾首屏速度和交互流畅度。"),
    ("5.10 有上限的字幕翻译缓存", "视频播放器在浏览器内用 targetLanguage + sourceText 组成翻译缓存键，并将缓存条目限制在 MAX_TRANSLATION_CACHE_ENTRIES=300。缓存命中时不重复调用翻译接口，超过上限时删除最早条目，避免长时间使用播放器导致内存无限增长。"),
    ("5.11 视频切换清晰度时保留播放状态", "播放器切换清晰度时记录当前播放时间和播放状态，替换 video source 后等待 loadedmetadata，再恢复 currentTime；如果切换前正在播放，则尝试继续播放。这样清晰度切换不会把用户强制跳回视频开头。"),
    ("5.12 地图坐标与瓦片预加载的本地化处理", "地图模块对办公室坐标进行 GCJ-02 转换后再交给高德瓦片，并提前预加载中心附近瓦片。代码不依赖运行时地理编码，使用固定坐标和固定标记，保证联系页地图结果稳定、可重复，并减少进入联系页后的空白等待。"),
    ("5.13 飞书事件同步的防抖调度", "收到连续的飞书变更事件时，服务端不是每个事件立即触发一次完整同步，而是通过 scheduleHonorCacheSync 和 scheduleNewsCacheSync 延迟合并事件，再执行一次同步。这能减少短时间内的重复 API 调用和缓存写入。"),
    ("5.14 新闻和荣誉的前台降级内容", "荣誉接口请求失败时，页面保留代码内置的证书卡片；新闻接口没有可用数据时，页面不会把已有结构强行替换成空白。动态数据与静态展示内容分离，使网络异常时官网仍保有基本的信息展示能力。"),
    ("5.15 视频 Range 响应支持拖动和断点播放", "Node 服务对视频请求解析 Range 请求头，返回 206 Partial Content、Content-Range、Accept-Ranges 和正确的 Content-Length。浏览器因此可以按片段读取视频，支持进度条拖动、断点续播和较大视频文件的渐进式播放。"),
]

for heading, body in sections:
    add_marked(heading, level=2)
    add_marked(body)

add_marked("本节总结：这些能力共同构成了网页的细节质量——滚动位置可恢复、异步数据不乱写 DOM、网络请求不重复、动画尊重用户偏好、图片和视频按需加载、外部同步避免重复执行。它们不是额外设想，而是当前代码已经实现的工程亮点。")

doc.save(OUTPUT)
print(OUTPUT)
