from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT='weimou_web——架构展示与集成讲解稿.docx'
d=Document(); s=d.sections[0]; s.top_margin=Inches(.7); s.bottom_margin=Inches(.7); s.left_margin=Inches(.8); s.right_margin=Inches(.8)
for n,z,c in [('Normal',10.5,'222222'),('Title',24,'163A5F'),('Heading 1',16,'163A5F'),('Heading 2',12,'2F6B8A')]:
    st=d.styles[n]; st.font.name='微软雅黑'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑'); st.font.size=Pt(z); st.font.color.rgb=RGBColor.from_string(c)
def shade(cell,fill):
    sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),fill); cell._tc.get_or_add_tcPr().append(sh)
def tbl(head,rows):
    t=d.add_table(rows=1,cols=len(head)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,x in enumerate(head): t.rows[0].cells[i].text=x; shade(t.rows[0].cells[i],'DCEAF3')
    for row in rows:
        for i,x in enumerate(row): t.add_row().cells[i].text=str(x)
    d.add_paragraph()
def code(x):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(.2); r=p.add_run(x); r.font.name='Consolas'; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string('1F4E79')

p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('weimou_web\n架构展示与集成讲解稿'); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=RGBColor(22,58,95)
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('可直接照读｜配合《代码架构与设计》表格使用').italic=True
d.add_paragraph('使用方式：先展示架构表，再按本文顺序打开代码文件。')
d.add_heading('一、开场',1)
d.add_paragraph('大家好，今天介绍微眸医疗官网的整体代码架构，以及页面、服务端、缓存和外部平台之间是如何集成的。这个项目采用 React 应用壳层、页面交互层、专项能力模块、Node.js 服务端和本地数据缓存的分层架构。')
d.add_paragraph('整体链路是：用户访问页面，进入 React 应用壳层；壳层识别 Hash 路由；页面层渲染内容并请求同源 API；Node.js 读取本地缓存或同步外部平台；最后把标准化数据返回给页面。')
d.add_heading('二、架构表怎么讲',1)
tbl(['层级','文件','职责'],[('应用壳层','src/main.jsx','React 入口、Hash 路由、页面生命周期'),('页面与交互层','app.js','页面模板、新闻、轮播、表单和 DOM 事件'),('国际化与样式','i18n.js / styles.css','中英文文案和视觉样式'),('专项能力','src/pot-player.js / src/leaflet-map.js','视频、字幕、地图等复杂功能'),('服务端','server.js','同源 API、校验、限流、外部平台调用'),('数据缓存','data/news / data/honors','新闻和荣誉数据的本地回退')])
d.add_paragraph('这张表不是按文件数量分类，而是按职责划分。每一层都有明确边界，修改时尽量只影响对应层。')
d.add_heading('三、React 入口和页面生命周期',1)
d.add_paragraph('打开 src/main.jsx。App 组件维护当前页面和语言。当 Hash 变化时更新 page 状态，页面随之切换。')
code('useEffect(() => {\n  const handleRoute = () => setPage(route());\n  window.addEventListener("hashchange", handleRoute);\n  return () => window.removeEventListener("hashchange", handleRoute);\n}, []);')
d.add_paragraph('这里值得借鉴的是：事件监听和清理成对出现。进入页面时初始化地图、视频和交互，离开页面时统一执行 cleanup，避免重复绑定和资源泄漏。')
code('return () => {\n  mapCleanup?.();\n  playerCleanups.forEach((cleanup) => cleanup());\n  interactionCleanup?.();\n};')
d.add_heading('四、页面如何集成 API',1)
d.add_paragraph('打开 app.js。新闻页面通过 /api/news 获取数据，浏览器不直接请求飞书等外部平台。前端只访问自己的同源接口。')
code('newsRequest = fetch("/api/news")\n  .then((response) => response.ok ? response.json() : null)\n  .catch(() => null)\n  .finally(() => { newsRequest = null; });')
d.add_paragraph('这里有两个设计点：前端短时间缓存，减少重复请求；请求去重，多个地方同时需要新闻时复用同一个 Promise。')
d.add_heading('五、Node.js 服务端和本地缓存',1)
d.add_paragraph('打开 server.js。新闻接口返回本地缓存，而不是让每一次页面访问都实时请求外部平台。')
code('async function handleNews(request, response) {\n  const cache = newsCache || readNewsCache();\n  sendJson(response, 200, {\n    news: cache?.news || [], source: "local-cache"\n  });\n}')
d.add_paragraph('服务端负责拉取远程数据、校验字段、下载图片并写入 data/news/news.json。页面读取和后台同步互相解耦，所以外部平台短暂异常时仍可展示上一次成功同步的数据。')
d.add_heading('六、表单安全边界',1)
d.add_paragraph('联系表单的链路是：前端提交 JSON，服务端检查请求类型、字段长度和格式，然后进行限流和幂等校验，最后才调用外部平台。')
code('const client = getClientAddress(request);\nconst rate = checkLeadRateLimit(client, Date.now());\nif (!rate.allowed) return sendJson(response, 429, { message: "提交过于频繁，请稍后再试" });')
code('const lead = await readJson(request, leadBodyMaxBytes);\nif (!lead || typeof lead !== "object") return sendJson(response, 400, { message: "提交内容无效" });')
d.add_paragraph('前端校验只负责提升体验，真正的安全校验、限流和重复提交保护必须放在服务端。')
d.add_heading('七、地图和视频的按需集成',1)
d.add_paragraph('地图和视频属于体积较大、生命周期复杂的功能，因此拆成独立模块。资源可以在浏览器空闲时预加载，但只有进入联系页时才初始化地图。')
code('export function preloadLeafletMapResources() {\n  return Promise.all([loadLeafletStyle(), loadLeaflet()]);\n}')
d.add_heading('八、最值得借鉴的实现',1)
tbl(['实现','解决的问题','可迁移场景'],[('生命周期 cleanup','避免事件重复绑定和资源泄漏','地图、视频、WebSocket、定时器'),('缓存优先读取','外部平台异常时页面仍可展示','新闻、荣誉、配置和字典'),('请求去重','避免同一数据被请求多次','列表、详情、搜索建议'),('服务端统一校验','防止绕过前端提交非法数据','表单、上传、Webhook'),('限流 + 幂等','防止刷接口和重复提交','表单、支付、预约')])
d.add_heading('九、现场演示顺序',1)
for x in ['展示架构表，说明六个层级和职责边界。','打开 src/main.jsx，演示 Hash 路由和 cleanup。','打开 app.js，演示 /api/news 请求和缓存。','打开 server.js，演示缓存返回、字段校验、限流和幂等。','打开 src/leaflet-map.js 或 src/pot-player.js，说明专项能力封装。','浏览器演示新闻页和联系页，最后说明外部平台异常时的回退策略。']: d.add_paragraph(x,style='List Number')
d.add_heading('十、结尾总结',1)
d.add_paragraph('这个项目的核心不是简单地把页面拆成几个文件，而是把页面展示、交互逻辑、专项能力、服务端接口和外部数据同步分别隔离。前端负责展示和交互，Node.js 负责统一接口与安全边界，本地缓存负责稳定性，专项模块负责封装地图和视频。')
d.add_paragraph('建议大家重点关注三类实现：生命周期清理、缓存优先和服务端统一校验。这三类设计也适用于大多数前端加服务端的业务系统。')
d.save(OUT); print(OUT)
